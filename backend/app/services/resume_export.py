import re
from weasyprint import HTML
from pathlib import Path

from docx import Document


EXPORT_DIR = Path("exports")


def ensure_export_dir() -> None:
    EXPORT_DIR.mkdir(parents=True, exist_ok=True)


def sanitize_filename(filename: str) -> str:
    cleaned = filename.strip().lower()
    cleaned = re.sub(r"[^a-zA-Z0-9_-]+", "_", cleaned)
    cleaned = cleaned.strip("_")

    if not cleaned:
        cleaned = "resume"

    return cleaned


def export_markdown_resume(
    filename: str,
    resume_markdown: str,
) -> dict:
    ensure_export_dir()

    safe_name = sanitize_filename(filename)
    final_filename = f"{safe_name}.md"
    file_path = EXPORT_DIR / final_filename

    file_path.write_text(resume_markdown, encoding="utf-8")

    return {
        "filename": final_filename,
        "file_path": str(file_path),
        "message": "Markdown resume exported successfully.",
    }


def export_html_resume(
    filename: str,
    resume_html: str,
) -> dict:
    ensure_export_dir()

    safe_name = sanitize_filename(filename)
    final_filename = f"{safe_name}.html"
    file_path = EXPORT_DIR / final_filename

    file_path.write_text(resume_html, encoding="utf-8")

    return {
        "filename": final_filename,
        "file_path": str(file_path),
        "message": "HTML resume exported successfully.",
    }


def add_markdown_line_to_doc(document: Document, line: str) -> None:
    stripped = line.strip()

    if not stripped:
        return

    if stripped.startswith("# "):
        document.add_heading(stripped[2:], level=0)

    elif stripped.startswith("## "):
        document.add_heading(stripped[3:], level=1)

    elif stripped.startswith("### "):
        document.add_heading(stripped[4:], level=2)

    elif stripped.startswith("- ") or stripped.startswith("* "):
        document.add_paragraph(stripped[2:], style="List Bullet")

    else:
        paragraph = document.add_paragraph()

        # Basic bold support for **text**
        parts = stripped.split("**")

        if len(parts) > 1:
            for index, part in enumerate(parts):
                run = paragraph.add_run(part)
                if index % 2 == 1:
                    run.bold = True
        else:
            paragraph.add_run(stripped)


def export_docx_resume(
    filename: str,
    resume_markdown: str,
) -> dict:
    ensure_export_dir()

    safe_name = sanitize_filename(filename)
    final_filename = f"{safe_name}.docx"
    file_path = EXPORT_DIR / final_filename

    document = Document()

    for line in resume_markdown.splitlines():
        add_markdown_line_to_doc(document, line)

    document.save(file_path)

    return {
        "filename": final_filename,
        "file_path": str(file_path),
        "message": "DOCX resume exported successfully.",
    }
    
def export_pdf_resume(
    filename: str,
    resume_html: str,
) -> dict:
    ensure_export_dir()

    safe_name = sanitize_filename(filename)
    final_filename = f"{safe_name}.pdf"
    file_path = EXPORT_DIR / final_filename

    HTML(string=resume_html).write_pdf(file_path)

    return {
        "filename": final_filename,
        "file_path": str(file_path),
        "message": "PDF resume exported successfully.",
    }
